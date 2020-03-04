import numpy as np
import csv
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

names = []
starters = []
mains = []
desserts = []
document = Document()
sections = document.sections
for section in sections:
    sections[-1].orientation = WD_ORIENT.LANDSCAPE
def read_file():
    with open('annual_dinner_menu_options.csv') as file:
        reader = csv.reader(file)
        for row in reader:
            temp_row = list(row)
            names.append(str(temp_row[1]))
            starters.append(str(temp_row[3]))
            mains.append(str(temp_row[4]))
            desserts.append(str(temp_row[5]))

def find_and_replace():
    for i in range(len(starters)):
        if "Salmon" in starters[i]:
            starters[i] = "London Cured Salmon"
        elif "Chicken" in starters[i]:
            starters[i] = "Oriental Style Chicken"
        elif "Crispy" in starters[i]:
            starters[i] = "Crispy Mozzarella Balls"

        if "Teriyaki" in mains[i]:
            mains[i] = "Teriyaki Glazed Duck"
        elif "Slow" in mains[i]:
            mains[i] = "Slow Braised Pork"
        elif "Butternut" in mains[i]:
            mains[i] = "Butternut Squash"

        if "Tiramisu" in desserts[i]:
            desserts[i] = "Tiramisu"

def write_name(i):
    style = document.styles['Normal']
    font = style.font
    font.name = 'Imprint MT Shadow'
    font.size = Pt(18)
    paragraph=document.add_paragraph(str(names[i]))
    paragraph.style = document.styles['Normal']
    sentence = paragraph.add_run("\n" + starters[i] + "\n" + mains[i] + "\n" + desserts[i])
    sentence.font.name = 'Avenir Next'
    sentence.font.size = Pt(12)
    #food = document.add_paragraph(starters[i] + "\n" + mains[1] + "\n" + desserts[1])


def write_food(i):
    style = document.styles['Normal']
    font = style.font
    font.name = 'Avenir Next'
    font.size = Pt(12)
    paragraph=document.add_paragraph("\n" + starters[i] + "\n" + mains[i] + "\n" + desserts[i] + "\n")
    paragraph.style = document.food_styles['Normal']

def create_table():
    table = document.add_table(rows = 220, cols = 3)
    table.autofit= False
    for cell in table.columns[0].cells:
        cell.width = Cm(8.32)
    for cell in table.columns[1].cells:
        cell.width = Cm(8.32)
    for cell in table.columns[2].cells:
        cell.width = Cm(8.32)
    for row in table.rows:
        row.height = Cm(4.42)
    table.style = 'Table Grid'
    return table

def add_to_table(table, i, row, column):
    mini = table.cell(row, column).add_table(rows=1, cols=2)
    mini.columns[0].cells[0].width = Cm(5)
    #mini.
    #mini.columns[0].cells[1].width = Cm(3.3)

    ### create the style
    style = document.styles['Normal']
    font = style.font
    font.name = 'Imprint MT Shadow'
    font.size = Pt(18)

    # paragraph = document.add_paragraph(str(names[i]))

    # sentence = paragraph.add_run("\n" + starters[i] + "\n" + mains[i] + "\n" + desserts[i])
    # sentence.font.name = 'Avenir Next'
    # sentence.font.size = Pt(12)
    paragraph = mini.cell(0, 0).add_paragraph(names[i], style='Normal')
    food = paragraph.add_run("\n" + starters[i] + "\n" + mains[i] + "\n" + desserts[i])
    food.font.name = 'Avenir Next'
    food.font.size = Pt(12)
    paragraph2 = mini.cell(0, 1).paragraphs[0]
    paragraph2 = mini.cell(0,1).add_paragraph("", style='Normal')
    run = paragraph2.add_run()
    run.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run.add_picture('pic.png', width=Cm(2.68))

def write_table(i):
    ### create the table
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    row_cells = table.add_row().cells
    mini = table.cell(0, 1).add_table(rows=1, cols=2)

    ### create the style
    style = document.styles['Normal']
    font = style.font
    font.name = 'Imprint MT Shadow'
    font.size = Pt(18)

    #paragraph = document.add_paragraph(str(names[i]))

    #sentence = paragraph.add_run("\n" + starters[i] + "\n" + mains[i] + "\n" + desserts[i])
    #sentence.font.name = 'Avenir Next'
    #sentence.font.size = Pt(12)
    paragraph = mini.cell(0, 0).add_paragraph(str(names[i]), style='Normal')
    food = paragraph.add_run("\n" + starters[i] + "\n" + mains[i] + "\n" + desserts[i])
    food.font.name = 'Avenir Next'
    food.font.size = Pt(12)
    paragraph2 = mini.cell(0, 1).paragraphs[0]
    run = paragraph2.add_run()
    run.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run.alignment = WD_ALIGN_PARAGRAPH.Low
    run.add_picture('pic.png', width=Cm(2.68))

    for cell in table.columns.cells:
        cell.width = Cm(8.32)
    for cell in table.columns[1].cells:
        cell.width = Cm(8.32)
    for cell in table.columns[2].cells:
        cell.width = Cm(8.32)
    for row in table.rows:
        row.height = Cm(4.42)
    mini.columns[0].cells[0].width = Cm(5)
    mini.columns[0].cells[1].width = Cm(3.3)
    table.style = 'Table Grid'

read_file()
find_and_replace()
#for i in range(1, len(mains)):
    #write_name(i)
#write_food(1)
#write_table(1)

table = create_table()
row = 1
for i in range(0, len(names)):
    column = (i%3)
    add_to_table(table, i, row, column)
    if i%3 == 2:
        row = row + 2
document.save("annual_dinner_test4.docx")