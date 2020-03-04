import numpy as np
import csv
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

names = []
starters = []
mains = []
desserts = []
club = []
document = Document()
sections = document.sections
for section in sections:
    sections[-1].orientation = WD_ORIENT.LANDSCAPE

def read_file():
    with open('springdin.csv') as file:
        reader = csv.reader(file)
        for row in reader:
            temp_row = list(row)
            names.append(str(temp_row[0]))
            starters.append(str(temp_row[1]))
            mains.append(str(temp_row[2]))
            desserts.append(str(temp_row[3]))
            club.append(temp_row[4])

def find_and_replace():
    for i in range(len(starters)):
        if "Salmon" in starters[i]:
            starters[i] = "Smoked Salmon"
        elif "Tarte Tatin" in starters[i]:
            starters[i] = "Tarte Tatin"

        if "Rib" in mains[i]:
            mains[i] = "Braised Rib of Beef"
        elif "Root" in mains[i]:
            mains[i] = "Vegetable Hot Pot"

def create_table():
    table = document.add_table(rows = 52, cols = 3)
    table.autofit= False
    for cell in table.columns[0].cells:
        cell.width = Cm(8.32)
    for cell in table.columns[1].cells:
        cell.width = Cm(8.32)
    for cell in table.columns[2].cells:
        cell.width = Cm(8.32)
    for row in table.rows:
        row.height = Cm(4.42)
    table.style = 'Table Grid'
    return table

def add_to_table(table, i, row, column):
    picture = table.cell((row-1),column).add_paragraph()
    picture_run = picture.add_run()
    picture_run.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_run.add_picture('XV22.png', width = Cm(7))
    mini = table.cell(row, column).add_table(rows=1, cols=2)
    mini.columns[0].cells[0].width = Cm(5)
    #mini.
    #mini.columns[0].cells[1].width = Cm(3.3)

    ### create the style
    style = document.styles['Normal']
    font = style.font
    font.name = 'Imprint MT Shadow'
    font.size = Pt(18)

    # paragraph = document.add_paragraph(str(names[i]))

    # sentence = paragraph.add_run("\n" + starters[i] + "\n" + mains[i] + "\n" + desserts[i])
    # sentence.font.name = 'Avenir Next'
    # sentence.font.size = Pt(12)
    paragraph = mini.cell(0, 0).add_paragraph(names[i], style='Normal')
    food = paragraph.add_run("\n" + starters[i] + "\n" + mains[i] + "\n" + desserts[i])
    food.font.name = 'Avenir Next'
    food.font.size = Pt(12)
    #paragraph2 = mini.cell(0, 1).paragraphs[0]
    paragraph2 = mini.cell(0,1).add_paragraph("", style='Normal')
    run = paragraph2.add_run()
    run.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if club[i] == '22':
        run.add_picture('22.png', width=Cm(2.68))
    else:
        run.add_picture('XV.png', width=Cm(2.68))

read_file()
find_and_replace()
table = create_table()
row = 1
for i in range(0, len(names)):
    column = (i%3)
    add_to_table(table, i, row, column)
    if i%3 == 2:
        row = row + 2
document.save("spring_dinner_test.docx")